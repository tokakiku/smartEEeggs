# ==========================================
# 灵犀智课 - Word 详案物理生成引擎 (命题终极优化版)
# 特性：严扣命题五大要素 (目标、过程、方法、活动、作业)
# ==========================================

import os
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger("Word_Engine")
OUTPUT_DIR = "downloads/exports"


def set_cell_background(cell, fill_color="EFEFEF"):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text_center(cell, text, bold=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def generate_word_from_json(json_data: dict, project_id: int) -> str:
    logger.info(f"📄 [详案引擎] 正在为项目 {project_id} 渲染命题级标准教案...")

    if "outline_data" in json_data and isinstance(json_data["outline_data"], dict):
        json_data = json_data["outline_data"]

    doc = Document()
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    metadata = json_data.get("course_metadata", {})
    syllabus = json_data.get("syllabus_content", [])

    # ==========================================
    # 0. 大标题
    # ==========================================
    heading = doc.add_heading('教 学 详 案', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()

    # ==========================================
    # 1. 教学基本信息 (新增【教学方法】)
    # ==========================================
    table_info = doc.add_table(rows=4, cols=4)
    table_info.style = 'Table Grid'

    cell_header1 = table_info.cell(0, 0)
    cell_header1.merge(table_info.cell(0, 3))
    set_cell_background(cell_header1)
    set_cell_text_center(cell_header1, "教学基本信息")

    table_info.cell(1, 0).text = "课题"
    set_cell_text_center(table_info.cell(1, 0), "课题")
    cell_title = table_info.cell(1, 1)
    cell_title.merge(table_info.cell(1, 3))
    set_cell_text_center(cell_title, metadata.get("title", "未命名课程"), bold=True)

    set_cell_text_center(table_info.cell(2, 0), "授课对象")
    set_cell_text_center(table_info.cell(2, 1), metadata.get("target_audience", "通用"))
    set_cell_text_center(table_info.cell(2, 2), "课程时长")
    set_cell_text_center(table_info.cell(2, 3), f"{metadata.get('total_duration', 45)} 分钟")

    # 🌟 命题点：教学方法提取
    set_cell_text_center(table_info.cell(3, 0), "教学方法")
    methods = metadata.get("teaching_methods", ["讲授法", "案例分析法"])
    method_str = "、".join(methods) if isinstance(methods, list) else str(methods)

    cell_method = table_info.cell(3, 1)
    cell_method.merge(table_info.cell(3, 3))
    set_cell_text_center(cell_method, method_str)

    doc.add_paragraph()

    # ==========================================
    # 2. 教学目标及重难点
    # ==========================================
    key_points, diff_points = [], []
    for item in syllabus:
        for kp in item.get("core_knowledge_points", []):
            if kp.get("is_key_point"): key_points.append(kp.get("point"))
            if kp.get("is_difficult_point"): diff_points.append(kp.get("point"))

    table_obj = doc.add_table(rows=2, cols=1)
    table_obj.style = 'Table Grid'
    set_cell_background(table_obj.cell(0, 0))
    set_cell_text_center(table_obj.cell(0, 0), "教学目标及教学重点、难点")

    p_obj = table_obj.cell(1, 0).paragraphs[0]
    p_obj.add_run("教学目标：\n").bold = True
    for i, obj in enumerate(metadata.get("teaching_objectives", [])):
        p_obj.add_run(f"  {i + 1}. {obj}\n")
    p_obj.add_run("\n重点：").bold = True
    p_obj.add_run("；".join(key_points) if key_points else "详见教学过程")
    p_obj.add_run("\n难点：").bold = True
    p_obj.add_run("；".join(diff_points) if diff_points else "详见教学过程")

    doc.add_paragraph()

    # ==========================================
    # 3. 教学过程核心表格 (凸显【课堂活动设计】)
    # ==========================================
    table_process = doc.add_table(rows=len(syllabus) + 2, cols=3)
    table_process.style = 'Table Grid'

    table_process.columns[0].width = Inches(1.0)
    table_process.columns[1].width = Inches(4.0)
    table_process.columns[2].width = Inches(1.5)

    cell_header3 = table_process.cell(0, 0)
    cell_header3.merge(table_process.cell(0, 2))
    set_cell_background(cell_header3)
    set_cell_text_center(cell_header3, "教学过程与课堂活动设计")  # 🌟 命题点：凸显表头

    set_cell_background(table_process.cell(1, 0))
    set_cell_text_center(table_process.cell(1, 0), "教学环节")
    set_cell_background(table_process.cell(1, 1))
    set_cell_text_center(table_process.cell(1, 1), "主要教学活动与内容")
    set_cell_background(table_process.cell(1, 2))
    set_cell_text_center(table_process.cell(1, 2), "意图及互动形式")

    for i, item in enumerate(syllabus):
        row = table_process.rows[i + 2].cells

        stage_name = item.get("stage", "")
        display_stage_name = stage_name.split("-")[-1].strip() if "-" in stage_name else stage_name
        set_cell_text_center(row[0], display_stage_name)

        activity_p = row[1].paragraphs[0]
        if "core_knowledge_points" in item and item["core_knowledge_points"]:
            for kp in item["core_knowledge_points"]:
                prefix = "🔥[难点] " if kp.get("is_difficult_point") else (
                    "⭐[重点] " if kp.get("is_key_point") else "🔸 ")
                run_title = activity_p.add_run(f"{prefix}{kp.get('point', '')}\n")
                run_title.bold = True
                if kp.get('explanation'):
                    activity_p.add_run(f"    {kp.get('explanation', '')}\n\n")
        else:
            activity_p.add_run(item.get("content_description", ""))

        intent_p = row[2].paragraphs[0]
        intent_p.add_run(
            f"【互动形式】\n{item.get('interaction_type', '讲授与演示')}\n\n【时间安排】\n{item.get('duration', 0)} 分钟")

    doc.add_paragraph()

    # ==========================================
    # 4. 课后作业模块 (新增【课后作业】)
    # ==========================================
    table_hw = doc.add_table(rows=2, cols=1)
    table_hw.style = 'Table Grid'

    set_cell_background(table_hw.cell(0, 0))
    set_cell_text_center(table_hw.cell(0, 0), "课后作业及拓展延伸")  # 🌟 命题点：作业模块

    hw_content = metadata.get("homework", "结合本节课重点，完成课后相关练习，并进行复盘总结。")
    hw_p = table_hw.cell(1, 0).paragraphs[0]
    hw_p.add_run(hw_content)

    # ==========================================
    # 5. 物理存储
    # ==========================================
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, f"word_project_{project_id}.docx")
    doc.save(output_path)

    logger.info(f"✅ [详案引擎] 命题级标准教案已生成: {output_path}")
    return output_path